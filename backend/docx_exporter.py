# -*- coding: utf-8 -*-
"""
docx_exporter.py — Xuất đề (định dạng \\begin{ex}...\\end{ex}) ra file Word .docx.

- Công thức Toán  -> OMML (công thức gốc của Word, sửa được). Cần file MML2OMML.XSL
  có sẵn khi đã cài Microsoft Office. Nếu không tìm thấy -> để nguyên chữ $...$.
- Hình TikZ       -> render PNG qua server compile-tikz-code.onrender.com rồi chèn ảnh.

Đường ống: LaTeX -> MathML (latex2mathml) -> OMML (lxml + MML2OMML.XSL).

Dùng chung cho app_math10/11/12 — đặt 1 bản trong mỗi thư mục lớp.
Hàm chính: export_questions_to_docx(full_text, output_path, meta=None, progress_cb=None)
"""

import os
import re
import glob
import copy
import time
import base64
from io import BytesIO
from html import escape as _html_escape
from contextvars import ContextVar

import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

try:
    from latex2mathml.converter import convert as _latex2mathml_convert
except Exception:
    _latex2mathml_convert = None

try:
    from PIL import Image, ImageChops
except Exception:
    Image = None
    ImageChops = None

# ----------------------------------------------------------------------------
# HẰNG SỐ
# ----------------------------------------------------------------------------
TIKZ_API_URL = "https://compile-tikz-code.onrender.com/compile"
M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

LABELS_MC = ["A", "B", "C", "D", "E", "F"]
LABELS_TF = ["a", "b", "c", "d", "e", "f"]
LABEL_BLUE = RGBColor(0x00, 0x00, 0xFF)       # tiền tố đáp án A./a)
CAU_VIOLET = RGBColor(0x8B, 0x5C, 0xF6)       # tiền tố "Câu X."
LOIGIAI_GREEN = RGBColor(0x00, 0x64, 0x00)    # chữ "Lời giải." (xanh lá đậm)

# Kiểu xuất công thức: "equation" -> OMML (sửa được trong Word),
#                     "latex"    -> giữ nguyên $..$ / $$..$$ (dán vào MathType).
_math_mode: ContextVar[str] = ContextVar("math_mode", default="equation")

# True: xuất đầy đủ (đánh dấu đáp án đúng + đáp số TLN + Lời giải).
# False: "chỉ có đề bài" — bỏ đánh dấu đáp án, bỏ [[đáp số]] TLN và bỏ Lời giải.
_show_answers: ContextVar[bool] = ContextVar("show_answers", default=True)

# Tiêu đề cho các nhóm câu hỏi (auto chèn theo qtype).
SECTION_TITLES = {
    "TN": "Phần I. TRẮC NGHIỆM",
    "DS": "Phần II. ĐÚNG SAI",
    "SA": "Phần III. TRẢ LỜI NGẮN",
    "TL": "Phần IV. TỰ LUẬN",
}
SECTION_ORDER = ["TN", "DS", "SA", "TL"]


def _style_label(run, correct):
    """Tiền tố đáp án: in đậm + màu xanh; gạch chân nếu là đáp án đúng."""
    run.bold = True
    run.font.color.rgb = LABEL_BLUE
    if correct:
        run.underline = True

# ----------------------------------------------------------------------------
# LATEX -> OMML
# ----------------------------------------------------------------------------
_xslt_cache = {"loaded": False, "transform": None}


def _find_mml2omml_xsl():
    """Tìm MML2OMML.XSL: (1) cạnh script (dùng khi deploy Linux); (2) vị trí Office trên
    Windows (dùng khi chạy máy giáo viên)."""
    here = os.path.dirname(os.path.abspath(__file__))
    local = os.path.join(here, "MML2OMML.XSL")
    if os.path.isfile(local):
        return local
    patterns = [
        r"C:\Program Files\Microsoft Office\root\Office*\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office*\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office*\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\Office*\MML2OMML.XSL",
    ]
    for pat in patterns:
        hits = glob.glob(pat)
        if hits:
            return hits[0]
    return None


def _get_omml_transform():
    if _xslt_cache["loaded"]:
        return _xslt_cache["transform"]
    _xslt_cache["loaded"] = True
    xsl = _find_mml2omml_xsl()
    if xsl:
        try:
            _xslt_cache["transform"] = etree.XSLT(etree.parse(xsl))
        except Exception:
            _xslt_cache["transform"] = None
    return _xslt_cache["transform"]


def omml_available():
    """True nếu chuyển được LaTeX -> OMML (đủ latex2mathml + MML2OMML.XSL)."""
    return _latex2mathml_convert is not None and _get_omml_transform() is not None


def _repair_omml_delims(omath):
    """Vá lỗi MML2OMML.XSL: với cặp ngoặc mà phần tử đầu bắt đầu bằng toán tử (vd dấu
    trừ trong '(-6;11)' hoặc '(-\\infty;..)'), XSL hiểu nhầm toán tử là 'sepChr' và tạo
    ô đối số RỖNG <m:e/> -> Word hiện thành ô vuông □. Sửa: gộp các <m:e>, đưa ký tự
    sep trở lại làm nội dung, bỏ sepChr. Chỉ động vào <m:d> CÓ <m:e> rỗng (đúng dấu hiệu lỗi)."""
    M = M_NS
    for d in list(omath.iter(M + "d")):
        dpr = d.find(M + "dPr")
        sep = dpr.find(M + "sepChr") if dpr is not None else None
        es = d.findall(M + "e")
        if sep is None or not any(len(e) == 0 for e in es):
            continue
        sepval = sep.get(M + "val") or sep.get("val") or ""
        new_e = etree.Element(M + "e")
        for idx, e in enumerate(es):
            if idx > 0 and sepval:
                r = etree.SubElement(new_e, M + "r")
                etree.SubElement(r, M + "t").text = sepval
            for child in list(e):
                new_e.append(child)
        for e in es:
            d.remove(e)
        dpr.remove(sep)
        d.append(new_e)
    return omath


def _latex_to_omath_elements(latex_src):
    """Trả về danh sách phần tử <m:oMath> hoặc None nếu thất bại."""
    if _latex2mathml_convert is None:
        return None
    transform = _get_omml_transform()
    if transform is None:
        return None
    try:
        mathml = _latex2mathml_convert(latex_src)
        mml_dom = etree.fromstring(mathml.encode("utf-8"))
        omml = transform(mml_dom)
        root = omml.getroot()
        if root is None:
            return None
        if root.tag == M_NS + "oMath":
            elems = [copy.deepcopy(root)]
        else:
            elems = [copy.deepcopy(e) for e in root.findall(".//" + M_NS + "oMath")]
        if not elems:
            return None
        for e in elems:
            _repair_omml_delims(e)
        return elems
    except Exception:
        return None


# ----------------------------------------------------------------------------
# TIỆN ÍCH PARSE LATEX (port từ texParser.ts của web)
# ----------------------------------------------------------------------------
def _extract_balanced(text, start_index):
    """Từ start_index quét tới dấu '{' đầu tiên, trả (nội_dung, vị_trí_'}')."""
    depth = 0
    content_start = -1
    for j in range(start_index, len(text)):
        c = text[j]
        if c == "{":
            if depth == 0:
                content_start = j + 1
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0 and content_start != -1:
                return text[content_start:j], j
    return None


def _skip_optional_brackets(text, i):
    while i < len(text) and text[i].isspace():
        i += 1
    if i < len(text) and text[i] == "[":
        while i < len(text) and text[i] != "]":
            i += 1
        return i + 1
    return i


def _convert_custom_ex_commands(text):
    """\\hoac{...} -> \\left[...\\right. ; \\heva{...} -> \\left\\{...\\right."""
    for cmd, prefix, suffix in (
        ("\\hoac", "\\left[\\begin{aligned}", "\\end{aligned}\\right."),
        ("\\heva", "\\left\\{\\begin{aligned}", "\\end{aligned}\\right."),
    ):
        idx = 0
        while True:
            pos = text.find(cmd, idx)
            if pos == -1:
                break
            res = _extract_balanced(text, pos + len(cmd))
            if not res:
                idx = pos + len(cmd)
                continue
            content, end = res
            repl = prefix + _convert_custom_ex_commands(content) + suffix
            text = text[:pos] + repl + text[end + 1:]
            idx = pos + len(repl)
    return text


def _unwrap_immini(text):
    """Mở phẳng \\immini[opt]{arg1}{arg2} -> arg1 + '\\n' + arg2."""
    out = ""
    last = 0
    pos = 0
    pat = re.compile(r"\\immini\s*(?:\[[^\]]*\])?\s*\{")
    while True:
        m = pat.search(text, pos)
        if not m:
            break
        out += text[last:m.start()]
        brace1 = m.end() - 1
        a1 = _extract_balanced(text, brace1)
        if not a1:
            out += m.group(0)
            last = m.end()
            pos = m.end()
            continue
        arg1, end1 = a1
        a2 = _extract_balanced(text, end1 + 1)
        if a2:
            arg2, end2 = a2
            out += arg1 + "\n" + arg2
            last = end2 + 1
        else:
            out += arg1
            last = end1 + 1
        pos = last
    out += text[last:]
    return out


def _strip_wrappers(text):
    text = re.sub(r"\\begin\{minipage\}(?:\{[^}]*\})*\s*", "", text)
    text = re.sub(r"\\end\{minipage\}\s*%?\s*", "", text)
    text = re.sub(r"\\centering\s*", "", text)
    text = re.sub(r"\\hspace\*?\{[^}]*\}\s*", "", text)
    text = re.sub(r"\\vspace\*?\{[^}]*\}\s*", "", text)
    return text


def _normalize_delims(text):
    return (
        text.replace("\\[", "$$").replace("\\]", "$$")
            .replace("\\(", "$").replace("\\)", "$")
    )


# Môi trường công thức "hiển thị" (đứng riêng, không cần $...$ trong LaTeX gốc).
_DISPLAY_ENV_RE = re.compile(
    r"\\begin\{(eqnarray\*?|align\*?|alignat\*?|gather\*?|multline\*?|equation\*?|displaymath)\}"
)

# Môi trường canh cột dùng '&' mà latex2mathml KHÔNG hiểu (xuất '&' sai XML).
# Phải đổi sang 'array' (latex2mathml xử lý '&' đúng) trước khi chuyển OMML.
_ALIGN_ENV_RE = re.compile(
    r"\\begin\{(aligned|alignedat|align\*?|alignat\*?|eqnarray\*?|split|cases|gathered)\}"
    r"(?:\{[^}]*\})?(?:\[[^\]]*\])?([\s\S]*?)\\end\{\1\}"
)


def _env_to_array(m):
    inner = m.group(2)
    rows = inner.split("\\\\")
    maxamp = max((r.count("&") for r in rows), default=0)
    cols = "l" * (maxamp + 1)
    return "\\begin{array}{" + cols + "}" + inner + "\\end{array}"


def _convert_math_envs(text):
    """Chuẩn hoá các môi trường công thức để latex2mathml chuyển được sang OMML."""
    # Bỏ "vỏ" môi trường math một dòng (nội dung bên trong đã là math).
    text = re.sub(r"\\begin\{(equation\*?|displaymath|gather\*?|multline\*?)\}", "", text)
    text = re.sub(r"\\end\{(equation\*?|displaymath|gather\*?|multline\*?)\}", "", text)
    # Đổi các môi trường canh cột -> array (lặp để xử lý lồng nhau).
    prev = None
    while prev != text:
        prev = text
        text = _ALIGN_ENV_RE.sub(_env_to_array, text)
    return text


# ----------------------------------------------------------------------------
# PHÂN TÍCH MỘT KHỐI ex
# ----------------------------------------------------------------------------
def _parse_choice_like(content, cmd):
    """Tách phần đề + 4 lựa chọn của \\choice / \\choiceTF."""
    idx = content.find(cmd)
    stem = content[:idx]
    cursor = idx + len(cmd)
    items = []
    for _ in range(6):
        b = content.find("{", cursor)
        if b == -1:
            break
        res = _extract_balanced(content, b)
        if not res:
            break
        raw, end = res
        correct = "\\True" in raw
        items.append({"text": raw.replace("\\True", " "), "correct": correct})
        cursor = end + 1
    return stem, items


def parse_ex(raw):
    """Phân tích nội dung 1 khối ex -> dict câu hỏi (giữ nguyên LaTeX thô)."""
    images = []

    def _repl(m):
        images.append(m.group(0))
        return " [IMG_ID_%d] " % (len(images) - 1)

    raw = re.sub(r"\\begin\s*\{tikzpicture\}[\s\S]*?\\end\s*\{tikzpicture\}", _repl, raw)
    # Bảng tabular -> render PNG (giống TikZ). Mảng 'array' để nguyên cho OMML.
    raw = re.sub(r"\\begin\s*\{tabular\}[\s\S]*?\\end\s*\{tabular\}", _repl, raw)
    flat = _strip_wrappers(_unwrap_immini(raw))
    flat = re.sub(r"(?<!\\)%.*", "", flat)   # bỏ comment LaTeX (vd nhãn %%[NB])

    explanation = ""
    content = flat
    lg = flat.find("\\loigiai")
    has_lg = lg != -1
    if has_lg:
        res = _extract_balanced(flat, lg)
        if res:
            explanation = res[0]
            content = flat[:lg]

    content = re.sub(r"\\diem\s*\{[^}]*\}", " ", content)

    # Phân loại DỰA HOÀN TOÀN vào mã lệnh trong khối ex (theo thứ tự ưu tiên):
    #   \choiceTF -> Đúng/Sai (DS), \shortans -> Trả lời ngắn (SA),
    #   \choice   -> Trắc nghiệm (TN), còn lại -> Tự luận (TL).
    # Lưu ý: "\choiceTF" chứa chuỗi con "\choice" nên has_choice phải loại trừ \choiceTF.
    has_tf = "\\choiceTF" in content
    has_sa = "\\shortans" in content
    has_choice = re.search(r"\\choice(?!TF)", content) is not None

    stem = content
    options = []
    statements = []
    answer = ""

    if has_tf:
        qtype = "DS"
        stem, statements = _parse_choice_like(content, "\\choiceTF")
    elif has_sa:
        qtype = "SA"
        idx = content.find("\\shortans")
        stem = content[:idx]
        astart = _skip_optional_brackets(content, idx + len("\\shortans"))
        b = content.find("{", astart)
        if b != -1:
            res = _extract_balanced(content, b)
            if res:
                answer = res[0].replace("\\True", " ")
    elif has_choice:
        qtype = "TN"
        stem, options = _parse_choice_like(content, "\\choice")
    else:
        # Không có mã lệnh nào -> câu tự luận (bất kể có \loigiai hay không).
        qtype = "TL"
        stem = content

    # Hậu kiểm: gắn lại placeholder hình bị sót (vd hình nằm sau \choice trong immini).
    # PHẢI tính cả explanation, nếu không ảnh trong lời giải sẽ bị gắn LẶP sang đề bài.
    captured = " ".join(
        [stem, explanation] + [o["text"] for o in options]
        + [s["text"] for s in statements] + [answer]
    )
    for k in range(len(images)):
        ph = "[IMG_ID_%d]" % k
        if ph not in captured:
            stem = stem + "\n" + ph

    return {
        "qtype": qtype,
        "stem": stem,
        "options": options,
        "statements": statements,
        "answer": answer,
        "explanation": explanation,
        "images": images,
    }


def parse_document(full_text):
    """Quét toàn văn -> danh sách item: {'type':'section'} hoặc {'type':'question'}."""
    items = []
    ex_pat = re.compile(r"\\begin\{ex\}([\s\S]*?)\\end\{ex\}")
    sec_pat = re.compile(r"\\section\*?\s*\{([^}]*)\}")
    last = 0
    for m in ex_pat.finditer(full_text):
        between = full_text[last:m.start()]
        for sm in sec_pat.finditer(between):
            items.append({"type": "section", "name": sm.group(1).strip()})
        q = parse_ex(m.group(1))
        items.append({"type": "question", "q": q})
        last = m.end()
    return items


# ----------------------------------------------------------------------------
# LÀM SẠCH ĐOẠN CHỮ THƯỜNG (ngoài công thức)
# ----------------------------------------------------------------------------
def _clean_plain(s):
    s = re.sub(r"(?<!\\)%.*", "", s)                       # comment
    s = s.replace("\\lq\\lq", '"').replace("\\rq\\rq", '"')
    # \dotfill / \hrulefill -> dòng chấm (PHẢI xử lý trước \dots vì \dotfill chứa \dots).
    s = re.sub(r"\\dotfill\b", "." * 40, s)
    s = re.sub(r"\\hrulefill\b", "." * 40, s)
    s = s.replace("\\dots", "…").replace("\\ldots", "…").replace("\\cdots", "…")
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    s = re.sub(r"\\includegraphics(?:\[[^\]]*\])?\{[^}]*\}", "", s)
    s = re.sub(r"\\(link|linkno)?ans\{[^}]*\}", "", s)
    s = re.sub(r"\\SA\s*(?:\[[^\]]*\])?\s*\{[^{}]*\}", "", s)   # ô đáp án trong đề
    s = re.sub(r"\\begin\{(itemize|enumerate|itemchoice|listEX|description|center)\}(?:\s*\[[^\]]*\])?", "", s)
    s = re.sub(r"\\end\{(itemize|enumerate|itemchoice|listEX|description|center)\}", "", s)
    s = re.sub(r"\\(itemch|item)\b", "\n• ", s)
    s = re.sub(r"\\(par|noindent|allowdisplaybreaks|notag|centering|displaystyle|limits|left|right)\b", " ", s)
    s = re.sub(r"\\(bf|it|rm|sf|tt|large|Large|small|normalsize|textnormal)\b", " ", s)
    s = s.replace("~", " ")
    s = re.sub(r"\\[,;:!> ]", " ", s)                      # lệnh khoảng cách
    s = s.replace("\\&", "&").replace("\\%", "%").replace("\\#", "#").replace("\\_", "_")
    s = s.replace("{", "").replace("}", "")
    s = s.replace("\t", " ")
    s = re.sub(r"[ ]{2,}", " ", s)
    s = re.sub(r"\n{2,}", "\n", s)
    return s


# ----------------------------------------------------------------------------
# GHÉP NỘI DUNG HỖN HỢP (chữ + công thức + hình) VÀO 1 PARAGRAPH
# ----------------------------------------------------------------------------
def _emit_text(paragraph, s, fmt):
    s = _clean_plain(s)
    if not s.strip():
        # vẫn giữ khoảng trắng đơn nếu cần ngăn cách
        if s and not s.isspace():
            return
    parts = s.split("\n")
    for k, part in enumerate(parts):
        if k > 0:
            paragraph.add_run().add_break()
        if part:
            r = paragraph.add_run(part)
            if fmt.get("bold"):
                r.bold = True
            if fmt.get("italic"):
                r.italic = True
            if fmt.get("underline"):
                r.underline = True


def _emit_math(paragraph, latex_math, fmt, display=False):
    latex_math = _convert_custom_ex_commands(latex_math.strip())
    if not latex_math.strip():
        return
    # Chế độ giữ nguyên LaTeX (để dán vào MathType): bọc $..$/$$..$$ và viết thẳng.
    if _math_mode.get() == "latex":
        delim = "$$" if display else "$"
        r = paragraph.add_run(delim + latex_math.strip() + delim)
        if fmt.get("bold"):
            r.bold = True
        if fmt.get("italic"):
            r.italic = True
        if fmt.get("underline"):
            r.underline = True
        return
    latex_math = _convert_math_envs(latex_math)
    elems = _latex_to_omath_elements(latex_math)
    if elems:
        for e in elems:
            paragraph._p.append(e)
    else:
        _emit_text(paragraph, "$" + latex_math + "$", fmt)


def _img_width_cm(png_bytes):
    if not Image:
        return None
    try:
        im = Image.open(BytesIO(png_bytes))
        w, _h = im.size
        return min(w / 300.0 * 2.54, 14.0)   # density=300 khi render
    except Exception:
        return None


def _trim_image(png_bytes):
    """Cắt sát nội dung (bỏ viền trống/nguyên trang). Server render tabular ra cả
    trang A4 -> phải cắt, nếu không khung ảnh cao bằng trang làm vỡ bố cục."""
    if not Image:
        return png_bytes
    try:
        im = Image.open(BytesIO(png_bytes)).convert("RGBA")
        bbox = im.getbbox()  # vùng không trong suốt
        if bbox is None or (
            bbox[2] - bbox[0] > im.width * 0.97 and bbox[3] - bbox[1] > im.height * 0.97
        ):
            # nền không trong suốt -> cắt theo viền trắng
            rgb = Image.new("RGB", im.size, (255, 255, 255))
            rgb.paste(im, mask=im.split()[3])
            diff = ImageChops.difference(rgb, Image.new("RGB", im.size, (255, 255, 255)))
            bbox = diff.getbbox()
        if bbox:
            pad = 6
            box = (max(0, bbox[0] - pad), max(0, bbox[1] - pad),
                   min(im.width, bbox[2] + pad), min(im.height, bbox[3] + pad))
            im = im.crop(box)
        out = BytesIO()
        im.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return png_bytes


def _emit_image(paragraph, entry):
    # entry: None | bytes | (png_bytes|None, reason) — chuẩn hoá về (bytes, reason).
    if isinstance(entry, tuple):
        png_bytes, reason = entry
    else:
        png_bytes, reason = entry, ""
    if not png_bytes:
        msg = " [Hình: không tải được%s] " % (" — " + reason if reason else "")
        _emit_text(paragraph, msg, {})
        return
    try:
        png_bytes = _trim_image(png_bytes)
        width = _img_width_cm(png_bytes)
        run = paragraph.add_run()
        if width:
            run.add_picture(BytesIO(png_bytes), width=Cm(width))
        else:
            run.add_picture(BytesIO(png_bytes))
    except Exception:
        _emit_text(paragraph, " [Hình: lỗi chèn] ", {})


def _add_inline(doc, paragraph, text, rendered, fmt=None, allow_blocks=False):
    """Phân tích text thành chữ/công thức/hình và ghi vào paragraph.
    allow_blocks=True (đề bài, lời giải): ảnh tách RA DÒNG RIÊNG, CANH GIỮA; văn bản
    sau ảnh viết tiếp ở paragraph mới. allow_blocks=False (đáp án, ô bảng): ảnh chèn
    nội dòng. Trả về paragraph cuối cùng đang dùng."""
    fmt = fmt or {}
    state = {"p": paragraph, "pending": False}

    def cur():
        # Tạo paragraph mới một cách "lười" để không sinh dòng trống thừa sau ảnh.
        if state["pending"]:
            state["p"] = doc.add_paragraph()
            state["pending"] = False
        return state["p"]

    text = _normalize_delims(_convert_custom_ex_commands(text.strip()))
    i, n = 0, len(text)
    buf = ""

    def flush():
        nonlocal buf
        if buf:
            _emit_text(cur(), buf, fmt)
            buf = ""

    while i < n:
        c = text[i]
        # công thức
        if c == "$":
            display = False
            if text[i:i + 2] == "$$":
                close = text.find("$$", i + 2)
                if close == -1:
                    buf += text[i:]
                    break
                math = text[i + 2:close]
                i = close + 2
                display = True
            else:
                close = text.find("$", i + 1)
                if close == -1:
                    buf += text[i:]
                    break
                math = text[i + 1:close]
                i = close + 1
            flush()
            _emit_math(cur(), math, fmt, display=display)
            continue
        # hình
        if c == "[" and text[i:i + 7] == "[IMG_ID":
            m = re.match(r"\[IMG_ID_(\d+)\]", text[i:])
            if m:
                flush()
                png = rendered.get(int(m.group(1)))
                if allow_blocks and doc is not None:
                    imgp = doc.add_paragraph()
                    imgp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _emit_image(imgp, png)
                    state["pending"] = True   # văn bản kế tiếp -> paragraph mới
                else:
                    _emit_image(cur(), png)
                i += m.end()
                continue
        # định dạng / xuống dòng / môi trường công thức hiển thị
        if c == "\\":
            me = _DISPLAY_ENV_RE.match(text, i)
            if me:
                endtok = "\\end{" + me.group(1) + "}"
                endpos = text.find(endtok, i)
                if endpos != -1:
                    flush()
                    _emit_math(cur(), text[i:endpos + len(endtok)], fmt, display=True)
                    i = endpos + len(endtok)
                    continue
            m = re.match(r"\\(textbf|textit|emph|underline|text)\s*\{", text[i:])
            if m:
                res = _extract_balanced(text, i + m.end() - 1)
                if res:
                    inner, end = res
                    flush()
                    nf = dict(fmt)
                    cmd = m.group(1)
                    if cmd == "textbf":
                        nf["bold"] = True
                    elif cmd in ("textit", "emph"):
                        nf["italic"] = True
                    elif cmd == "underline":
                        nf["underline"] = True
                    _add_inline(doc, cur(), inner, rendered, nf, allow_blocks=False)
                    i = end + 1
                    continue
            if text[i:i + 2] == "\\\\":
                flush()
                cur().add_run().add_break()
                i += 2
                m2 = re.match(r"\s*\[[^\]]*\]", text[i:])
                if m2:
                    i += m2.end()
                continue
        buf += c
        i += 1
    flush()
    return state["p"]


# ----------------------------------------------------------------------------
# RENDER TIKZ
# ----------------------------------------------------------------------------
def compile_tikz_b64(tikz_code, transparent=True, timeout=120, retries=7, on_retry=None):
    """Gọi server TikZ để biên dịch mã -> PNG. Trả (image_base64, reason, is_compile_error).

    - image_base64: chuỗi base64 của PNG nếu thành công, None nếu hỏng.
    - reason: '' khi OK; ngược lại là mô tả ngắn lý do hỏng.
    - is_compile_error: True nếu server đã chạy nhưng MÃ TIKZ SAI (thử lại vô ích) ->
      caller nên báo 422; False nếu lỗi tạm thời do server ngủ/mạng -> nên báo 502.
    - on_retry(attempt, reason): gọi TRƯỚC mỗi lần chờ thử lại (để caller phát nhịp tim
      giữ kết nối stream sống, tránh proxy Render cắt khi chờ server ngủ dậy).

    Server compile-tikz (Render gói free) hay ngủ: lần gọi đầu thường trả 502/503 kèm
    body HTML (không phải JSON) hoặc timeout. Cold-start có thể mất ~30-60s, nên ta THỬ
    LẠI KIÊN NHẪN (mặc định 7 lần, tổng giãn cách ~63s) để "sống lâu hơn" thời gian server
    thức dậy — giống trình duyệt bên app exam chờ không timeout. CHỈ khi server trả JSON
    hợp lệ mà thiếu ảnh mới coi là lỗi mã (dừng ngay, không thử lại)."""
    last = "không rõ"
    for attempt in range(retries):
        try:
            resp = requests.post(
                TIKZ_API_URL,
                json={
                    "source": tikz_code,
                    "mode": "auto",
                    "format": "png",
                    "density": 300,
                    "transparent": bool(transparent),
                    "return_log": True,
                },
                timeout=timeout,
            )
        except requests.exceptions.Timeout:
            last = "server phản hồi chậm (có thể đang khởi động)"
        except requests.exceptions.RequestException:
            last = "không gọi được server"
        else:
            # Gateway của Render khi server đang thức dậy -> body không phải JSON -> thử lại.
            if resp.status_code in (502, 503, 504):
                last = "server đang khởi động"
            else:
                try:
                    data = resp.json()
                except ValueError:
                    # 2xx nhưng body rỗng/không phải JSON (proxy cold-start) -> thử lại.
                    last = "server phản hồi rỗng"
                else:
                    if resp.ok and data.get("image_base64"):
                        return data["image_base64"], "", False
                    # Server chạy bình thường, trả JSON nhưng không có ảnh -> MÃ TIKZ SAI.
                    log = (data.get("log") or data.get("error") or "").strip()
                    msg = "mã TikZ biên dịch lỗi" + ((": " + log[:1500]) if log else "")
                    return None, msg, True
        # Còn lượt -> chờ giãn cách tăng dần rồi thử lại (đánh thức server ngủ).
        if attempt < retries - 1:
            if on_retry:
                on_retry(attempt + 1, last)
            time.sleep(3 * (attempt + 1))
    return None, last, False


def _render_tikz(tikz_code, timeout=120, retries=7, on_retry=None):
    """Bọc compile_tikz_b64 cho luồng xuất Word: trả (png_bytes, reason)."""
    b64, reason, _is_compile = compile_tikz_b64(
        tikz_code, transparent=True, timeout=timeout, retries=retries, on_retry=on_retry
    )
    if b64:
        try:
            return base64.b64decode(b64), ""
        except Exception:
            return None, "ảnh trả về hỏng"
    return None, reason


# ----------------------------------------------------------------------------
# DỰNG TÀI LIỆU
# ----------------------------------------------------------------------------
def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")


def _add_header(doc, meta, made_override=None):
    school = (meta.get("school") or "").strip()
    title = (meta.get("title") or "").strip()
    subject = (meta.get("subject") or "").strip()
    duration = (meta.get("duration") or "").strip()
    klass = (meta.get("class_name") or "").strip()
    made = (made_override if made_override is not None else (meta.get("made") or "")).strip()

    # Tiêu đề xếp trong 1 bảng 2 cột (không viền) — như bố cục đề LaTeX.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.width = Cm(7.5)
    right.width = Cm(8.5)

    def line(cell, text, bold=False, italic=False, size=None, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)

    # Cột TRÁI: Trường / Mã đề
    fl = True
    if school:
        line(left, school, bold=True, first=True); fl = False
    if made:
        line(left, "Mã đề: " + made, bold=True, first=fl); fl = False
    if fl:
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cột PHẢI: Tiêu đề / Môn / Thời gian
    fr = True
    if title:
        line(right, title, bold=True, size=13, first=True); fr = False
    if subject:
        line(right, subject, bold=True, first=fr); fr = False
    if duration:
        line(right, "Thời gian: " + duration, first=fr); fr = False
    if fr:
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dòng họ tên (toàn chiều ngang, dưới bảng)
    doc.add_paragraph()
    namep = doc.add_paragraph()
    namep.add_run(
        "Họ và tên thí sinh: " + "." * 30 + "  Lớp: " + (klass or "." * 6)
        + "   Điểm: " + "." * 6
    )
    doc.add_paragraph()


def _opt_len(text):
    """Ước lượng độ dài hiển thị 1 lựa chọn (bỏ ký hiệu $)."""
    return len(_clean_plain(text.replace("$", "")).strip())


def _add_mc_options(doc, options, rendered):
    if not options:
        return
    show = _show_answers.get()
    maxlen = max((_opt_len(o["text"]) for o in options), default=0)
    per_row = 4 if maxlen <= 12 else 2 if maxlen <= 30 else 1
    n = len(options)

    if per_row == 1:
        for idx, o in enumerate(options):
            mark = show and o["correct"]
            p = doc.add_paragraph()
            _style_label(p.add_run("%s. " % LABELS_MC[idx]), mark)
            _add_inline(doc, p, o["text"], rendered, {"underline": mark})
        return

    rows = (n + per_row - 1) // per_row
    table = doc.add_table(rows=rows, cols=per_row)
    table.autofit = True
    for idx, o in enumerate(options):
        mark = show and o["correct"]
        r, c = idx // per_row, idx % per_row
        para = table.cell(r, c).paragraphs[0]
        _style_label(para.add_run("%s. " % LABELS_MC[idx]), mark)
        _add_inline(doc, para, o["text"], rendered, {"underline": mark})


def _add_tf_statements(doc, statements, rendered):
    show = _show_answers.get()
    for idx, s in enumerate(statements):
        mark = show and s["correct"]
        p = doc.add_paragraph()
        _style_label(p.add_run("%s) " % LABELS_TF[idx]), mark)
        _add_inline(doc, p, s["text"], rendered, {"underline": mark})


def _add_question(doc, qno, q, rendered):
    qtype = q["qtype"]
    p = doc.add_paragraph()
    cr = p.add_run("Câu %d. " % qno)
    cr.bold = True
    cr.font.color.rgb = CAU_VIOLET
    # ĐS (đúng/sai): đánh dấu [TF] tô đậm ngay sau tiền tố "Câu X." (quy ước exam-genius-app).
    if qtype == "DS":
        p.add_run("[TF] ").bold = True
    _add_inline(doc, p, q["stem"], rendered, {}, allow_blocks=True)

    if qtype == "TN":
        _add_mc_options(doc, q["options"], rendered)
    elif qtype == "DS":
        _add_tf_statements(doc, q["statements"], rendered)
    elif qtype == "SA" and _show_answers.get():
        # TLN: đáp số trong cặp [[...]] đặt ở DÒNG RIÊNG, tách khỏi giả thiết.
        ap = doc.add_paragraph()
        ap.add_run("[[")
        _add_inline(doc, ap, q["answer"], rendered, {})
        ap.add_run("]]")
    # TL: phần hỏi đã nằm trong stem, không thêm gì.

    if _show_answers.get() and q["explanation"].strip():
        lab = doc.add_paragraph()
        lab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lab.add_run("Lời giải.")
        r.bold = True
        r.font.color.rgb = LOIGIAI_GREEN
        body = doc.add_paragraph()
        _add_inline(doc, body, q["explanation"], rendered, {}, allow_blocks=True)


def _emit_section_title(doc, title):
    """Chèn 1 tiêu đề mục ('Phần I. TRẮC NGHIỆM'...) — canh giữa, đậm, tím."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = CAU_VIOLET


def export_questions_to_docx(
    full_text, output_path, meta=None, progress_cb=None, math_mode="equation",
    show_header=True, show_answers=True,
):
    """
    Phân tích full_text (toàn văn LaTeX có các khối ex) và ghi ra file .docx.
    - meta: {school, class_name, title, subject, duration, made}
    - progress_cb(done, total, message) để báo tiến trình render hình (tùy chọn).
    - math_mode: "equation" (mặc định, OMML — sửa được trong Word) hoặc
                 "latex" (giữ nguyên $..$ / $$..$$ để dán vào MathType).
    - show_header: True -> in phần tiêu đề (trường/mã đề/họ tên) VÀ các đề mục
                   "Phần I/II/III/IV"; False -> chỉ in các câu hỏi, bỏ hết tiêu đề/đề mục.
    - show_answers: True -> đầy đủ (đánh dấu đáp án đúng + đáp số TLN + Lời giải);
                    False -> "chỉ có đề bài" (bỏ đánh dấu đáp án, đáp số và lời giải).
    Câu hỏi tự gom theo loại và (khi show_header) chèn tiêu đề "Phần I/II/III/IV..."
    (chỉ chèn các phần thực sự có câu). Số câu đánh liên tục xuyên qua các phần.
    Trả về output_path.
    """
    meta = meta or {}
    base_made = str(meta.get("made") or "").strip()

    token = _math_mode.set(math_mode if math_mode in ("equation", "latex") else "equation")
    ans_token = _show_answers.set(bool(show_answers))
    try:
        # Tách thành nhiều ĐỀ theo mốc "% Đề N" mà worker chèn giữa các đề.
        # Mỗi đề lặp lại header; mã đề tăng dần từng đơn vị (+1).
        chunks = [c for c in re.split(r"%\s*Đề\s*\d+\s*", full_text) if "\\begin{ex}" in c]
        if not chunks:
            chunks = [full_text]

        parsed = [parse_document(c) for c in chunks]
        total_imgs = sum(
            len(it["q"]["images"]) for items in parsed for it in items if it["type"] == "question"
        )
        done_imgs = 0

        doc = Document()
        _set_base_style(doc)

        for di, items in enumerate(parsed):
            if di > 0:
                doc.add_page_break()
            made_i = str(int(base_made) + di) if base_made.isdigit() else base_made
            if show_header:
                _add_header(doc, meta, made_override=made_i)

            # Gom câu hỏi theo qtype để chèn tiêu đề mục "Phần I/II/III/IV".
            # Các \section{} có sẵn trong nguồn được bỏ qua ở chế độ auto này.
            by_type = {"TN": [], "DS": [], "SA": [], "TL": []}
            for it in items:
                if it["type"] == "question":
                    by_type.setdefault(it["q"]["qtype"], []).append(it["q"])

            qno = 0
            for qtype in SECTION_ORDER:
                qs = by_type.get(qtype) or []
                if not qs:
                    continue
                if show_header:
                    _emit_section_title(doc, SECTION_TITLES[qtype])
                for q in qs:
                    qno += 1
                    rendered = {}
                    for k, code in enumerate(q["images"]):
                        if progress_cb:
                            progress_cb(done_imgs, total_imgs,
                                        "Đang vẽ hình %d/%d" % (done_imgs + 1, max(1, total_imgs)))
                        # Nhịp tim khi server hình còn ngủ: giữ kết nối stream sống,
                        # tránh Render cắt trong lúc chờ ~60s cold-start.
                        def _hb(attempt, reason, _d=done_imgs):
                            if progress_cb:
                                progress_cb(_d, total_imgs,
                                            "Máy chủ hình đang khởi động, thử lại lần %d…" % attempt)
                        rendered[k] = _render_tikz(code, on_retry=_hb)
                        done_imgs += 1
                    _add_question(doc, qno, q, rendered)

        doc.save(output_path)
        return output_path
    finally:
        _math_mode.reset(token)
        _show_answers.reset(ans_token)


# ----------------------------------------------------------------------------
# XUẤT HTML (để in ra PDF ở trình duyệt)
# ----------------------------------------------------------------------------
# Dùng lại TOÀN BỘ parser + render TikZ ở trên; chỉ đổi "đầu ra": thay vì ghi vào
# python-docx, ta ghép chuỗi HTML. Công thức giữ nguyên LaTeX và để MathJax (nạp qua
# CDN trong tài liệu) render; hình TikZ nhúng base64. Bố cục/màu bám theo bản .docx.

def _esc(s):
    return _html_escape(s, quote=False)


def _text_html(s, fmt):
    """Đoạn chữ thường -> HTML (đã làm sạch lệnh LaTeX, escape, xuống dòng -> <br>)."""
    s = _clean_plain(s)
    parts = s.split("\n")
    chunks = []
    for k, part in enumerate(parts):
        if k > 0:
            chunks.append("<br>")
        if part:
            chunks.append(_esc(part))
    inner = "".join(chunks)
    if not inner.strip():
        return inner  # giữ khoảng trắng ngăn cách nếu có
    if fmt.get("bold"):
        inner = "<b>%s</b>" % inner
    if fmt.get("italic"):
        inner = "<i>%s</i>" % inner
    if fmt.get("underline"):
        inner = "<u>%s</u>" % inner
    return inner


def _math_html(latex_math, display=False):
    """Công thức -> để MathJax render: bọc \\(..\\) (inline) hoặc \\[..\\] (display).
    Escape '<' và '&' để trình duyệt không hiểu nhầm là thẻ; MathJax vẫn đọc đúng ký tự."""
    latex_math = _convert_custom_ex_commands(latex_math.strip())
    if not latex_math.strip():
        return ""
    body = latex_math.replace("&", "&amp;").replace("<", "&lt;")
    if display:
        return '<div class="math-display">\\[%s\\]</div>' % body
    return "\\(%s\\)" % body


def _img_html(entry, block=False):
    """Placeholder hình -> thẻ <img> base64 (hoặc thông báo lỗi nếu render hỏng)."""
    if isinstance(entry, tuple):
        png_bytes, reason = entry
    else:
        png_bytes, reason = entry, ""
    if not png_bytes:
        msg = "[Hình: không tải được%s]" % (" — " + reason if reason else "")
        return '<span class="img-fail">%s</span>' % _esc(msg)
    try:
        png_bytes = _trim_image(png_bytes)
        width = _img_width_cm(png_bytes)
        b64 = base64.b64encode(png_bytes).decode("ascii")
        style = ' style="width:%.2fcm"' % width if width else ""
        img = '<img class="q-img" src="data:image/png;base64,%s"%s>' % (b64, style)
    except Exception:
        img = '<span class="img-fail">[Hình: lỗi chèn]</span>'
    if block:
        return '<div class="img-center">%s</div>' % img
    return img


def _inline_html(text, rendered, fmt=None, block_img=False):
    """Song song với _add_inline nhưng sinh chuỗi HTML thay vì ghi paragraph.
    block_img=True (đề bài, lời giải): ảnh xuống DÒNG RIÊNG, canh giữa."""
    fmt = fmt or {}
    text = _normalize_delims(_convert_custom_ex_commands(text.strip()))
    i, n = 0, len(text)
    buf = ""
    out = []

    def flush():
        nonlocal buf
        if buf:
            out.append(_text_html(buf, fmt))
            buf = ""

    while i < n:
        c = text[i]
        # công thức
        if c == "$":
            display = False
            if text[i:i + 2] == "$$":
                close = text.find("$$", i + 2)
                if close == -1:
                    buf += text[i:]
                    break
                math = text[i + 2:close]
                i = close + 2
                display = True
            else:
                close = text.find("$", i + 1)
                if close == -1:
                    buf += text[i:]
                    break
                math = text[i + 1:close]
                i = close + 1
            flush()
            out.append(_math_html(math, display))
            continue
        # hình
        if c == "[" and text[i:i + 7] == "[IMG_ID":
            m = re.match(r"\[IMG_ID_(\d+)\]", text[i:])
            if m:
                flush()
                out.append(_img_html(rendered.get(int(m.group(1))), block=block_img))
                i += m.end()
                continue
        # định dạng / xuống dòng / môi trường công thức hiển thị
        if c == "\\":
            me = _DISPLAY_ENV_RE.match(text, i)
            if me:
                endtok = "\\end{" + me.group(1) + "}"
                endpos = text.find(endtok, i)
                if endpos != -1:
                    flush()
                    out.append(_math_html(text[i:endpos + len(endtok)], display=True))
                    i = endpos + len(endtok)
                    continue
            m = re.match(r"\\(textbf|textit|emph|underline|text)\s*\{", text[i:])
            if m:
                res = _extract_balanced(text, i + m.end() - 1)
                if res:
                    inner, end = res
                    flush()
                    nf = dict(fmt)
                    cmd = m.group(1)
                    if cmd == "textbf":
                        nf["bold"] = True
                    elif cmd in ("textit", "emph"):
                        nf["italic"] = True
                    elif cmd == "underline":
                        nf["underline"] = True
                    out.append(_inline_html(inner, rendered, nf, block_img=False))
                    i = end + 1
                    continue
            if text[i:i + 2] == "\\\\":
                flush()
                out.append("<br>")
                i += 2
                m2 = re.match(r"\s*\[[^\]]*\]", text[i:])
                if m2:
                    i += m2.end()
                continue
        buf += c
        i += 1
    flush()
    return "".join(out)


def _mc_options_html(options, rendered):
    if not options:
        return ""
    show = _show_answers.get()
    maxlen = max((_opt_len(o["text"]) for o in options), default=0)
    per_row = 4 if maxlen <= 12 else 2 if maxlen <= 30 else 1
    cells = []
    for idx, o in enumerate(options):
        mark = show and o["correct"]
        cls = "opt-label correct" if mark else "opt-label"
        lab = '<span class="%s">%s.</span> ' % (cls, LABELS_MC[idx])
        body = _inline_html(o["text"], rendered, {"underline": bool(mark)})
        cells.append('<td class="opt-cell">%s%s</td>' % (lab, body))
    rows = []
    for r in range(0, len(cells), per_row):
        rows.append("<tr>%s</tr>" % "".join(cells[r:r + per_row]))
    return '<table class="opt-table"><tbody>%s</tbody></table>' % "".join(rows)


def _tf_statements_html(statements, rendered):
    show = _show_answers.get()
    out = []
    for idx, s in enumerate(statements):
        mark = show and s["correct"]
        cls = "opt-label correct" if mark else "opt-label"
        lab = '<span class="%s">%s)</span> ' % (cls, LABELS_TF[idx])
        body = _inline_html(s["text"], rendered, {"underline": bool(mark)})
        out.append('<div class="tf-item">%s%s</div>' % (lab, body))
    return "".join(out)


def _question_html(qno, q, rendered):
    qtype = q["qtype"]
    prefix = '<span class="cau">Câu %d.</span> ' % qno
    if qtype == "DS":
        prefix += "<b>[TF]</b> "
    parts = ['<p class="stem">%s%s</p>'
             % (prefix, _inline_html(q["stem"], rendered, {}, block_img=True))]

    if qtype == "TN":
        parts.append(_mc_options_html(q["options"], rendered))
    elif qtype == "DS":
        parts.append(_tf_statements_html(q["statements"], rendered))
    elif qtype == "SA" and _show_answers.get():
        ans = _inline_html(q["answer"], rendered, {})
        parts.append('<p class="sa-ans">[[%s]]</p>' % ans)

    if _show_answers.get() and q["explanation"].strip():
        parts.append('<p class="loigiai-label">Lời giải.</p>')
        parts.append('<div class="loigiai">%s</div>'
                     % _inline_html(q["explanation"], rendered, {}, block_img=True))
    return '<div class="question">%s</div>' % "".join(parts)


def _header_html(meta, made_override=None):
    school = _esc((meta.get("school") or "").strip())
    title = _esc((meta.get("title") or "").strip())
    subject = _esc((meta.get("subject") or "").strip())
    duration = _esc((meta.get("duration") or "").strip())
    klass = _esc((meta.get("class_name") or "").strip())
    made = made_override if made_override is not None else (meta.get("made") or "")
    made = _esc(str(made).strip())

    left = []
    if school:
        left.append('<p class="hdr-line"><b>%s</b></p>' % school)
    if made:
        left.append('<p class="hdr-line"><b>Mã đề: %s</b></p>' % made)

    right = []
    if title:
        right.append('<p class="hdr-line hdr-title"><b>%s</b></p>' % title)
    if subject:
        right.append('<p class="hdr-line"><b>%s</b></p>' % subject)
    if duration:
        right.append('<p class="hdr-line">Thời gian: %s</p>' % duration)

    name_line = ("Họ và tên thí sinh: " + "." * 30 + "  Lớp: "
                 + (klass or "." * 6) + "   Điểm: " + "." * 6)
    return (
        '<table class="hdr-table"><tbody><tr>'
        '<td class="hdr-left">%s</td><td class="hdr-right">%s</td>'
        '</tr></tbody></table>'
        '<p class="name-line">%s</p>' % ("".join(left), "".join(right), name_line)
    )


_HTML_CSS = """
* { box-sizing: border-box; }
body { font-family: "Times New Roman", serif; font-size: 12pt; color: #000;
       margin: 0; line-height: 1.35; }
.page { padding: 1.2cm 1.4cm; }
.page + .page { page-break-before: always; }
.hdr-table { width: 100%; border-collapse: collapse; margin-bottom: 6pt; }
.hdr-table td { vertical-align: top; text-align: center; }
.hdr-left { width: 45%; } .hdr-right { width: 55%; }
.hdr-line { margin: 0; text-align: center; }
.hdr-title { font-size: 13pt; }
.name-line { margin: 10pt 0 14pt; }
.section-title { text-align: center; font-weight: bold; font-size: 13pt;
                 color: #8B5CF6; margin: 12pt 0 6pt; }
.question { margin: 0 0 9pt; }
.stem { margin: 4pt 0; text-align: justify; }
.cau { font-weight: bold; color: #8B5CF6; }
.opt-table { border-collapse: collapse; margin: 2pt 0; width: 100%; }
.opt-cell { vertical-align: top; padding: 1pt 8pt 1pt 0; }
.opt-label { font-weight: bold; color: #0000FF; }
.opt-label.correct { text-decoration: underline; }
.tf-item { margin: 2pt 0; }
.sa-ans { margin: 2pt 0; }
.loigiai-label { text-align: center; font-weight: bold; color: #006400;
                 margin: 5pt 0 2pt; }
.loigiai { margin: 0 0 6pt; text-align: justify; }
.img-center { text-align: center; margin: 5pt 0; }
.q-img { max-width: 100%; vertical-align: middle; }
.img-fail { color: #b00020; font-style: italic; }
/* Chống tách khối khi phân trang: cả câu, ô đáp án, hình… không bị cắt ngang qua 2 trang.
   Với câu quá dài vượt 1 trang thì trình duyệt vẫn buộc phải ngắt, nhưng ngắt theo DÒNG
   nên không xẻ ngang chữ. */
.question { page-break-inside: avoid; break-inside: avoid; }
.opt-table, .opt-cell, .tf-item, .sa-ans, .img-center, .stem { page-break-inside: avoid;
                                                               break-inside: avoid; }
.section-title, .loigiai-label { page-break-after: avoid; break-after: avoid; }
@page { margin: 1.2cm 1.4cm; }
@media print { .page { padding: 0; } }
"""

# MathJax + tự mở hộp thoại In sau khi render xong (người dùng chọn "Lưu thành PDF").
# Dùng bản SVG (không phải CHTML), fontCache='none': công thức thành hình vector TỰ CHỨA,
# in ra chuẩn, KHÔNG phụ thuộc web-font — tránh lỗi công thức rỗng khi Chrome in trước lúc
# nạp font. Trình duyệt phân trang theo DÒNG (không cắt ngang chữ) và chữ vẫn chọn/copy được.
_HTML_HEAD_SCRIPT = """
<script>
window.MathJax = {
  tex: { inlineMath: [['\\\\(', '\\\\)']], displayMath: [['\\\\[', '\\\\]'], ['$$', '$$']] },
  svg: { fontCache: 'none' },
  startup: {
    pageReady: function () {
      return MathJax.startup.defaultPageReady().then(function () {
        setTimeout(function () { window.focus(); window.print(); }, 300);
      });
    }
  }
};
</script>
<script async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-svg.js"></script>
"""


def export_questions_to_html(
    full_text, meta=None, progress_cb=None, show_header=True, show_answers=True,
    pdf_name="de_thi.pdf",
):
    """Như export_questions_to_docx nhưng TRẢ VỀ chuỗi HTML (một tài liệu hoàn chỉnh có
    MathJax + html2pdf) để dựng PDF và TẢI THẲNG về ở trình duyệt. Dùng lại parser +
    render TikZ; tôn trọng show_header (tiêu đề + đề mục) và show_answers (đầy đủ / chỉ
    đề bài). pdf_name: tên file .pdf tải về."""
    meta = meta or {}
    base_made = str(meta.get("made") or "").strip()
    ans_token = _show_answers.set(bool(show_answers))
    try:
        chunks = [c for c in re.split(r"%\s*Đề\s*\d+\s*", full_text) if "\\begin{ex}" in c]
        if not chunks:
            chunks = [full_text]

        parsed = [parse_document(c) for c in chunks]
        total_imgs = sum(
            len(it["q"]["images"]) for items in parsed for it in items if it["type"] == "question"
        )
        done_imgs = 0

        pages = []
        for di, items in enumerate(parsed):
            made_i = str(int(base_made) + di) if base_made.isdigit() else base_made
            body = []
            if show_header:
                body.append(_header_html(meta, made_override=made_i))

            by_type = {"TN": [], "DS": [], "SA": [], "TL": []}
            for it in items:
                if it["type"] == "question":
                    by_type.setdefault(it["q"]["qtype"], []).append(it["q"])

            qno = 0
            for qtype in SECTION_ORDER:
                qs = by_type.get(qtype) or []
                if not qs:
                    continue
                if show_header:
                    body.append('<div class="section-title">%s</div>'
                                % _esc(SECTION_TITLES[qtype]))
                for q in qs:
                    qno += 1
                    rendered = {}
                    for k, code in enumerate(q["images"]):
                        if progress_cb:
                            progress_cb(done_imgs, total_imgs,
                                        "Đang vẽ hình %d/%d" % (done_imgs + 1, max(1, total_imgs)))
                        # Nhịp tim khi server hình còn ngủ: giữ kết nối stream sống,
                        # tránh Render cắt trong lúc chờ ~60s cold-start.
                        def _hb(attempt, reason, _d=done_imgs):
                            if progress_cb:
                                progress_cb(_d, total_imgs,
                                            "Máy chủ hình đang khởi động, thử lại lần %d…" % attempt)
                        rendered[k] = _render_tikz(code, on_retry=_hb)
                        done_imgs += 1
                    body.append(_question_html(qno, q, rendered))

            pages.append('<div class="page">%s</div>' % "".join(body))

        # <title> = tên file (bỏ .pdf): Chrome dùng nó làm tên gợi ý khi "Lưu thành PDF".
        doc_title = re.sub(r"\.pdf$", "", pdf_name, flags=re.I).strip() or "de_thi"
        return (
            "<!doctype html>\n<html lang=\"vi\"><head><meta charset=\"utf-8\">"
            "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
            "<title>%s</title>\n<style>%s</style>\n%s</head>"
            "<body>%s</body></html>"
            % (_esc(doc_title), _HTML_CSS, _HTML_HEAD_SCRIPT, "".join(pages))
        )
    finally:
        _show_answers.reset(ans_token)
